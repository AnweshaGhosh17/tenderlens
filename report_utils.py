from docx import Document
from docx.shared import Pt
from datetime import datetime

def generate_docx_report(report_text, output_path):
    doc = Document()

    # Style configuration (Optional: makes the font look cleaner)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('TenderLens AI Analysis Report', level=0)
    
    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated on: {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    date_run.italic = True
    
    # Process lines
    lines = report_text.split("\n")
    
    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            continue

        # Intelligent Heading Detection
        if clean_line.startswith("###"):
            doc.add_heading(clean_line.replace("###", "").strip(), level=2)
        elif clean_line.startswith("##"):
            doc.add_heading(clean_line.replace("##", "").strip(), level=1)
        elif clean_line.endswith(":") and len(clean_line) < 40:
            # Short lines ending in colon are usually section headers
            doc.add_heading(clean_line.replace(":", ""), level=3)
        elif clean_line.startswith("- ") or clean_line.startswith("* "):
            # Handle Bullet points properly
            doc.add_paragraph(clean_line[2:], style='List Bullet')
        else:
            # Normal paragraph text, clean up any remaining bold markers
            text = clean_line.replace("**", "")
            doc.add_paragraph(text)

    # Footer - Using a proper footer section instead of just a page break
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = f"TenderLens AI Internal Report - {datetime.now().year}"

    doc.save(output_path)
